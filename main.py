from flask import Flask, request, jsonify
from flask_cors import CORS
from google.cloud import storage, language_v1
import os
import PyPDF2
import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import tempfile
from datetime import datetime
from werkzeug.utils import secure_filename

app = Flask(__name__)
CORS(app)

UPLOAD_BUCKET = 'resume-uploads-2024'
OUTPUT_BUCKET = 'resume-outputs-2024'

os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "GOOGLE_APPLICATION_CREDENTIALS=C:\path\to\your\key.json"

def extract_text_from_pdf(file_stream):
    reader = PyPDF2.PdfReader(file_stream)
    return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

def extract_text_from_docx(file_stream):
    doc = docx.Document(file_stream)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def analyze_text_with_nlp(text):
    client = language_v1.LanguageServiceClient()
    document = language_v1.Document(content=text, type_=language_v1.Document.Type.PLAIN_TEXT)
    response = client.analyze_entities(request={'document': document})
    keywords = [entity.name for entity in response.entities]
    return keywords

def generate_resume_pdf(text, filename):
    temp_path = os.path.join(tempfile.gettempdir(), filename)
    c = canvas.Canvas(temp_path, pagesize=letter)
    c.drawString(100, 750, "AI-Enhanced Resume")
    c.drawString(100, 735, f"Generated on {datetime.now().strftime('%Y-%m-%d')}")
    c.drawString(100, 700, "Content Preview:")
    y = 680
    for line in text.splitlines():
        if y < 100: break
        c.drawString(100, y, line[:100])
        y -= 15
    c.save()
    return temp_path

def upload_to_bucket(bucket_name, source_file_name, dest_blob_name):
    client = storage.Client()
    bucket = client.get_bucket(bucket_name)
    blob = bucket.blob(dest_blob_name)
    blob.upload_from_filename(source_file_name)

@app.route('/upload', methods=['POST'])
def upload_resume():
    if request.content_type.startswith('multipart/form-data'):
        if 'resume' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['resume']
        filename = secure_filename(file.filename)

        if filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file)
        elif filename.lower().endswith('.docx'):
            text = extract_text_from_docx(file)
        else:
            return jsonify({'error': 'Only PDF or DOCX resumes are supported'}), 400

        upload_to_bucket(UPLOAD_BUCKET, file.stream.name, filename)

    elif request.content_type == 'application/pdf':
        temp_path = os.path.join(tempfile.gettempdir(), 'raw_resume.pdf')
        with open(temp_path, 'wb') as f:
            f.write(request.get_data())
        with open(temp_path, 'rb') as f:
            text = extract_text_from_pdf(f)
        filename = 'raw_resume.pdf'
        upload_to_bucket(UPLOAD_BUCKET, temp_path, filename)

    elif request.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        temp_path = os.path.join(tempfile.gettempdir(), 'raw_resume.docx')
        with open(temp_path, 'wb') as f:
            f.write(request.get_data())
        with open(temp_path, 'rb') as f:
            text = extract_text_from_docx(f)
        filename = 'raw_resume.docx'
        upload_to_bucket(UPLOAD_BUCKET, temp_path, filename)

    else:
        return jsonify({'error': 'Unsupported file format or content type'}), 400

    keywords = analyze_text_with_nlp(text)
    enhanced_text = f"{text}\n\nSuggested Keywords to Include:\n" + ", ".join(keywords)

    output_filename = filename.replace('.pdf', '_enhanced.pdf').replace('.docx', '_enhanced.pdf')
    output_path = generate_resume_pdf(enhanced_text, output_filename)
    upload_to_bucket(OUTPUT_BUCKET, output_path, output_filename)

    return jsonify({
        'message': 'Resume uploaded and processed',
        'download_url': f'https://storage.googleapis.com/{OUTPUT_BUCKET}/{output_filename}'
    })

if __name__ == '__main__':
    app.run(debug=True)
